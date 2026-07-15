from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = r"C:\Users\mizuno\Documents\製図のAIドリブン化\AIサポート製図_検討整理.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "667085"
WHITE = "FFFFFF"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:' + m))
        if node is None:
            node = OxmlElement('w:' + m)
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i]))
            tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_run_font(run, size=11, color=None, bold=None, italic=None):
    run.font.name = 'Yu Gothic'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Yu Gothic')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Yu Gothic')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def style_paragraph(p, before=0, after=6, line=1.1):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line

def add_para(doc, text='', size=11, color='000000', bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, before, after, 1.1)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size, color, bold, italic)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, 16, BLUE, True)
    elif level == 2:
        set_run_font(r, 13, BLUE, True)
    else:
        set_run_font(r, 12, DARK, True)
    return p

def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, widths)
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_fill)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, 10, INK, True)
    for rowdata in rows:
        cells = table.add_row().cells
        for i, val in enumerate(rowdata):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            set_run_font(r, 10)
    return table

def add_callout(doc, label, text, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [9360])
    cell = table.cell(0,0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + '  ')
    set_run_font(r, 10.5, DARK, True)
    r2 = p.add_run(text)
    set_run_font(r2, 10.5)
    return table

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Yu Gothic'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Yu Gothic')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Yu Gothic')
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
normal.font.size = Pt(11)
for name, size, color, before, after in [
    ('Heading 1',16,BLUE,16,8), ('Heading 2',13,BLUE,12,6), ('Heading 3',12,DARK,8,4)]:
    st = styles[name]
    st.font.name = 'Yu Gothic'
    st._element.rPr.rFonts.set(qn('w:ascii'), 'Yu Gothic')
    st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Yu Gothic')
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Header/footer
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run('AIサポート製図プロジェクト｜検討整理')
set_run_font(hr, 8.5, MUTED)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run('社内検討用｜2026年7月14日')
set_run_font(fr, 8.5, MUTED)

# Title block
add_para(doc, '検討整理', 10, BLUE, True, before=8, after=4)
add_para(doc, 'AIサポート体制でスピードアップを目指す製図・設計フロー', 23, INK, True, after=5)
add_para(doc, 'Tfas／DXF／AI／JW-CADを起点とした段階的な検証と、将来のIFC・BIM対応に向けた方向性', 11.5, MUTED, after=14)

meta = add_table(doc, ['項目','内容'], [
    ('目的','製図工程の短縮と、将来のBIM移行に対応できるAI活用基盤の検討'),
    ('対象業務','Tfasで共有された3D情報をもとに、社内ダクト加工用図面を作成・出力する業務'),
    ('現時点の位置づけ','検証開始前の論点整理・仮説・進め方のたたき台'),
], [1800, 7560])

add_heading(doc, '1. 目指す主な業務フロー', 1)
add_para(doc, 'Tfasで共有された3D情報を起点に、社内ダクト加工用の図面をDXFで出力し、JW-CADで開いて現場作業に使える状態まで整える。AIは最終成果物を無条件に確定するのではなく、図面化・確認・修正を支援する役割として位置づける。')

flow = add_table(doc, ['段階','入力／処理','出力・確認ポイント'], [
    ('1｜情報受領','Tfasで共有された3D情報を受領','対象範囲、属性、欠損情報を確認'),
    ('2｜図面化','社内ダクト加工向けのルールに沿って図面化','加工に必要な寸法・形状・接続情報を反映'),
    ('3｜DXF出力','加工用図面をDXF形式で出力','レイヤー、尺度、線種、座標、文字を確認'),
    ('4｜AI支援','複数パターンの候補や修正案を生成','候補の比較、採用案の選択、差分確認'),
    ('5｜JW-CAD確認','DXFをJW-CADで開く','現場で扱える表示・編集状態かを確認'),
    ('6｜現場利用','必要な修正を反映して作業に使用','作業結果・修正内容を次回検証へフィードバック'),
], [1500, 3600, 4260], header_fill=LIGHT_BLUE)

add_callout(doc, '基本方針', 'AIの出力を一つの正解として扱わず、複数候補を提示し、そのうち少なくとも一つを実務のたたき台として利用できる状態を目指す。')

add_heading(doc, '2. 関係するツール・サービス', 1)
tools = add_table(doc, ['ツール／サービス','現状・役割','検討事項'], [
    ('Tfas','開発用アカウント作成済み。3D情報の起点。','データ取得方法、出力形式、属性情報、検証用サンプルの確保'),
    ('JW-CAD','無料で利用可能。現場作業用の確認・編集環境。','DXFの互換性、表示崩れ、編集性、現場で必要な最低限の品質基準'),
    ('Cloud Run','ERP社内主幹システムの一部が稼働開始中。','専用プロジェクトを新設するか、入口・認証を共通化するかを検討。入口は共通、実行基盤・権限・課金単位は分離する案が有力'),
    ('AI API','目的に適したAPIが提供されているか未調査。','図面・幾何情報・属性情報を扱えること、候補生成、構造化出力、監査ログ、データ取り扱いを確認'),
    ('Rebro／BIM','業界で採用が進み、BIM利用が受注条件となる可能性。','将来のIFC入出力やBIM属性をツール側で吸収できる仕様にする'),
], [1750, 3560, 4050], header_fill=LIGHT_BLUE)

add_heading(doc, '3. まず確認すべきゴール', 1)
add_para(doc, '最初に、短期の改善と中長期の変革を同じ指標で評価しないよう、ゴールを二層に分けて定義する。')
add_table(doc, ['ゴールの層','定義案','評価指標の例'], [
    ('短期｜工程改善','製図の一部工程を10〜20％でも早くし、やり直しを増やさない','作業時間、修正回数、候補採用率、現場での手戻り件数'),
    ('中長期｜移行準備','BIM利用が求められた際に、データや運用を大きく作り直さず移行できる体制を整える','IFC対応範囲、属性の保持率、他ツール連携、再利用可能なデータ構造'),
], [1800, 4200, 3360])
add_callout(doc, '確認したい意思決定', '当面は「10〜20％の工程短縮」を優先するのか、それとも「BIM移行に備えたデータ・システム基盤」を優先するのか。両方を目指す場合も、短期検証の合格条件を先に定める。', fill='FFF8E8')

add_heading(doc, '4. 現時点で把握している利用者の感触', 1)
add_bullet(doc, '「やり直しになる可能性があるなら、AIを使わず最初から自分でやる」という意見が強い。')
add_bullet(doc, '一つの高精度な回答を待つより、複数パターンの候補から、どれかをたたき台として使える状態のほうが受け入れられやすい。')
add_bullet(doc, '製図工程は個人差が大きいため、平均値だけでなく、担当者ごとの作業時間・修正傾向・得意不得意を把握する必要がある。')
add_para(doc, 'このため、検証では「AIの正解率」だけでなく、候補が出るまでの時間、候補を採用・修正するまでの時間、最終的な手戻りの有無を測定する。')

add_heading(doc, '5. 検証の進め方', 1)
add_heading(doc, '5-1. 先行検証：Tfas → DXF → AI → JW-CAD', 2)
add_para(doc, '既存のTfas・DXF・JW-CADの流れを維持しながら、AIを補助工程として挿入し、導入効果とリスクを短期間で確認する。')
for s in [
    '代表的な図面・部材・形状を選び、検証用サンプルを準備する。',
    '現行手順で作成した正解図面と、作業時間・修正箇所を記録する。',
    'AIに候補生成・図面要素の抽出・チェック支援を行わせ、複数候補を比較する。',
    'DXFをJW-CADで開き、表示・編集・現場利用に問題がないか確認する。',
    '採用率、修正時間、手戻り、担当者の受容性を評価し、次の対象範囲を決める。',
]: add_number(doc, s)

add_heading(doc, '5-2. 将来検証：Tfas → IFC → BIM → AI → JW-CAD', 2)
add_para(doc, 'BIMが受注条件となる可能性を見据え、IFCを中間形式として扱えるか、属性情報が維持されるかを検証する。')
add_bullet(doc, 'TfasからIFCへ変換した際に、形状・寸法・接続・部材属性がどこまで保持されるか確認する。')
add_bullet(doc, 'BIM側で追加・補正された情報を、AIが読み取りやすい構造化データとして扱えるか確認する。')
add_bullet(doc, 'BIMで得た情報から、現場向けDXF/JW-CAD成果物へ戻す際の変換ルールを整理する。')
add_bullet(doc, '将来ツールがRebroなどに変わっても、データ形式・属性・変換ルールを吸収できる設計にする。')

add_heading(doc, '6. システム構成に関する検討事項', 1)
add_heading(doc, '6-1. Cloud Runプロジェクトの扱い', 2)
add_para(doc, 'ERP社内主幹システムとは別に、製図・AI支援用の専用プロジェクトを新設する案を基本の検討候補とする。ただし、利用者の入口は共通化し、認証・アカウント管理の負担を抑える方向が考えられる。')
add_table(doc, ['論点','共通化するもの','分離を検討するもの'], [
    ('利用者体験','入口、ログイン、基本的なアカウント情報','製図プロジェクト固有の画面・操作'),
    ('セキュリティ','認証基盤、利用者の所属情報','サービスアカウント、API権限、データアクセス範囲'),
    ('運用・費用','監視方針、問い合わせ窓口','Cloud Run、ストレージ、AI APIの課金・上限管理'),
    ('障害影響','共通ログインの障害対応','製図・AI側の障害がERPへ波及しない構成'),
], [1700, 3700, 3960])
add_callout(doc, '暫定的な見立て', '入口は同じでも、製図・AI支援の実行環境と権限は専用プロジェクトとして分離するほうが、検証・費用管理・将来拡張を進めやすい可能性がある。最終判断は、認証方式、データ連携、運用体制を確認して決定する。')

add_heading(doc, '7. 検証前に決める項目', 1)
for s in [
    '対象とする図面・部材・形状の範囲を限定する。',
    '現行作業の基準時間と、許容する修正回数・手戻りを定義する。',
    'AI出力を採用する人、最終確認する人、現場へ展開する人を決める。',
    '「採用可能な候補」と判断する品質基準を決める。',
    'DXF／IFCに含めるべき属性・レイヤー・尺度・命名規則を整理する。',
    '図面データをAIサービスへ送る場合の機密性・保存期間・ログ管理を確認する。',
    'Tfas、Rebro、JW-CAD、AI APIそれぞれの利用条件・ライセンス・サポート範囲を確認する。',
]: add_bullet(doc, s)

add_heading(doc, '8. 当面のアクション案', 1)
add_table(doc, ['順番','アクション','成果物・判断'], [
    ('1','検証対象を1〜2種類の代表図面に絞る','対象図面、入力データ、現行成果物'),
    ('2','現行作業を計測する','工程別時間、修正箇所、担当者差'),
    ('3','Tfas→DXF→JW-CADの互換性を確認する','変換ルールと品質チェック項目'),
    ('4','AI API候補を調査し、候補生成の小さな試作を行う','候補出力、採用率、修正時間'),
    ('5','Cloud Runの専用プロジェクト案と認証連携案を比較する','構成案、権限案、費用・運用論点'),
    ('6','IFC・BIM対応の必要属性を洗い出す','将来拡張を阻害しないデータ仕様'),
], [900, 4380, 4080])

add_heading(doc, '9. まとめ', 1)
add_para(doc, '本プロジェクトは、まず既存のTfas→DXF→JW-CADの業務を対象に、AIを候補生成・確認・修正支援として組み込み、実際の手戻りを増やさず工程を短縮できるか検証する。そのうえで、IFC・BIM・Rebro等への対応を見据え、ツールを固定するのではなく、データ形式・属性・変換ルールを吸収できる構成へ段階的に広げる。')
add_callout(doc, '次に決めること', '短期検証の対象図面、現行基準時間、合格条件、検証責任者を決定する。')

doc.core_properties.title = 'AIサポート体制でスピードアップを目指す製図・設計フロー'
doc.core_properties.subject = 'Tfas／DXF／AI／JW-CADを起点とした検討整理'
doc.core_properties.author = '社内検討用'
doc.save(OUT)
print(OUT)
